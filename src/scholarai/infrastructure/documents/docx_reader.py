"""DOCX text extraction via python-docx."""

from __future__ import annotations

import io

import docx

from scholarai.domain.models.documents import Document, DocumentType
from scholarai.infrastructure.documents.classification import classify_document
from scholarai.infrastructure.observability import get_logger

log = get_logger(__name__)


class DocxDocumentReader:
    def supports(self, filename: str) -> bool:
        return filename.lower().endswith(".docx")

    def read(self, filename: str, content: bytes) -> Document:
        try:
            document = docx.Document(io.BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
        except Exception as exc:  # noqa: BLE001 - malformed docx must degrade, not crash
            log.warning("documents.docx_reader.failed", filename=filename, error=str(exc))
            return Document(filename=filename, document_type=DocumentType.UNREADABLE, readable=False)

        if not text:
            return Document(filename=filename, document_type=DocumentType.UNREADABLE, readable=False)
        return Document(
            filename=filename,
            document_type=classify_document(filename, text),
            raw_text=text,
            page_count=1,
            readable=True,
        )
